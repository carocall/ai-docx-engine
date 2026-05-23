"""
表格块处理器（严格模式）
"""
from docx.shared import Pt

from .base import BlockHandler
from ..parser import TableBlock
from ..utils.borders import BorderHelper
from ..utils.spacing import SpacingHelper
from ..styles import StyleNotFoundError


class TableHandler(BlockHandler):
    """表格块处理器"""

    def can_handle(self, block) -> bool:
        return isinstance(block, TableBlock)

    def handle(self, block: TableBlock):
        """处理表格块"""
        if block.rows <= 0 or block.cols <= 0:
            return None

        # 验证表头样式
        if block.header_style:
            self.style_engine.require_style(block.header_style, "table header")
        else:
            raise StyleNotFoundError("table: 未指定 header_style")

        # 验证表体样式
        if block.body_style:
            self.style_engine.require_style(block.body_style, "table body")
        else:
            raise StyleNotFoundError("table: 未指定 body_style")

        # 创建表格
        table = self.doc.add_table(rows=block.rows, cols=block.cols)

        # 填充单元格
        for i in range(min(block.rows, len(block.cells))):
            row_cells = block.cells[i]
            for j in range(min(block.cols, len(row_cells))):
                cell = table.cell(i, j)
                cell.text = str(row_cells[j])

                # 应用单元格样式
                cell_style_name = block.header_style if i == 0 else block.body_style

                for para in cell.paragraphs:
                    para.style = cell_style_name
                    para.paragraph_format.left_indent = 0
                    para.paragraph_format.first_line_indent = 0
                    para.paragraph_format.hanging_indent = 0

        # 应用边框
        BorderHelper.apply_table_borders(table, block.border)

        # 表格后的间距
        para = self.doc.add_paragraph()
        SpacingHelper.set_spacing(para.paragraph_format, "space_after", block.space_after)

        return table
