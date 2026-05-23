"""
文本块处理器（严格模式）
"""
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .base import BlockHandler
from ..parser import TextBlock
from ..styles import StyleNotFoundError


class TextHandler(BlockHandler):
    """文本块处理器"""

    def can_handle(self, block) -> bool:
        return isinstance(block, TextBlock)

    def handle(self, block: TextBlock):
        """处理文本块"""
        # 确定样式名称
        if block.style:
            # 用户指定了样式，严格检查必须存在
            style_name = block.style
            self.style_engine.require_style(style_name, "text")
        else:
            # 未指定样式，报错
            raise StyleNotFoundError("text: 未指定样式名称")

        # 创建段落
        para = self.doc.add_paragraph(style=style_name)

        # 处理 runs
        if not block.runs:
            # 如果没有 runs，添加空段落
            return para

        for run_data in block.runs:
            self._add_run(para, run_data)

        return para

    def _add_run(self, para, run_data: dict):
        """添加run到段落"""
        text = run_data.get("text", "")
        if not text:
            return

        r = para.add_run(text)

        # 应用内联样式
        if run_data.get("bold"):
            r.bold = True
        if run_data.get("italic"):
            r.italic = True
        if run_data.get("underline"):
            r.underline = True
        if run_data.get("superscript"):
            r.font.superscript = True
        if run_data.get("subscript"):
            r.font.subscript = True

        # 文字颜色
        if run_data.get("color"):
            try:
                color = run_data["color"]
                if isinstance(color, str) and color.startswith("#"):
                    color = color[1:]
                r.font.color.rgb = RGBColor.from_string(color)
            except:
                pass

        # 高亮背景
        if run_data.get("highlight"):
            try:
                highlight_color = run_data["highlight"]
                if isinstance(highlight_color, str) and highlight_color.startswith("#"):
                    highlight_color = highlight_color[1:]
                rPr = r._element.get_or_add_rPr()
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), highlight_color)
                rPr.append(shading)
            except:
                pass
