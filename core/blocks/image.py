"""
图片块处理器（严格模式）
"""
from pathlib import Path
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .base import BlockHandler
from ..parser import ImageBlock
from ..utils.spacing import SpacingHelper
from ..styles import StyleNotFoundError


class ImageHandler(BlockHandler):
    """图片块处理器"""

    def can_handle(self, block) -> bool:
        return isinstance(block, ImageBlock)

    def handle(self, block: ImageBlock):
        """处理图片块"""
        # 确定样式
        if block.style:
            # 用户指定了样式，严格检查必须存在
            style_name = block.style
            self.style_engine.require_style(style_name, "image")
            style = self.style_engine.get_style(style_name)
        else:
            # 未指定样式，报错
            raise StyleNotFoundError("image: 未指定样式名称")

        # 创建段落
        para = self.doc.add_paragraph()

        # 设置对齐
        alignment = style.get("alignment", "center")
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT
        }
        para.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)

        # 设置间距
        SpacingHelper.set_spacing(para.paragraph_format, "space_before", style.get("space_before"))
        SpacingHelper.set_spacing(para.paragraph_format, "space_after", style.get("space_after"))

        # 处理图片路径
        if not block.src:
            para.add_run("[图片路径为空]")
            return para

        img_path = block.src
        if not Path(block.src).is_absolute():
            img_path = str(Path(self.base_dir) / block.src)

        # 插入图片
        try:
            width_cm = style.get("width_cm", 10)
            para.add_run().add_picture(img_path, width=Cm(width_cm))
        except Exception as e:
            para.add_run(f"[图片加载失败: {block.src}]")

        return para
