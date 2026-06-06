"""
标题块处理器 - 自动设置 outlineLevel（严格模式）
"""
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .base import BlockHandler
from ..parser import HeadingBlock
from ..styles import StyleNotFoundError


class HeadingHandler(BlockHandler):
    """标题块处理器"""

    def can_handle(self, block) -> bool:
        return isinstance(block, HeadingBlock)

    def handle(self, block: HeadingBlock):
        """处理标题块"""
        level = max(1, min(block.level, 4))  # 限制在 1-4

        # 样式名称必须显式指定
        if not block.style:
            raise StyleNotFoundError(f"heading (level={level}): 未指定 style")

        style_name = block.style
        self.style_engine.require_style(style_name, f"heading (level={level})")

        # 创建段落并应用样式
        para = self.doc.add_paragraph(style=style_name)

        # 设置 outlineLevel：优先从样式读取，否则从 block.level 推导
        style_props = self.style_engine.get_style(style_name)
        if style_props.get("outline_level") is not None:
            # 样式已定义 outline_level，注册时已生效，无需重复设置
            pass
        else:
            # 回退到 block.level
            self._set_outline_level(para, level)

        # 设置书签 id（如果提供了）
        if block.id:
            self._set_bookmark(para, block.id)

        # 处理 runs
        for run_data in block.runs:
            self._add_run(para, run_data)

        return para

    def _set_outline_level(self, para, level: int):
        """设置段落的 outlineLevel，使 Word 识别为标题"""
        pPr = para._element.get_or_add_pPr()
        outlineLvl = OxmlElement('w:outlineLvl')
        outlineLvl.set(qn('w:val'), str(level - 1))  # Word 中 0-based
        pPr.append(outlineLvl)

    def _set_bookmark(self, para, bookmark_id: str):
        """设置书签（用于内部跳转）"""
        # 在段落开始处插入书签
        run = para.runs[0] if para.runs else para.add_run()
        tag_start = OxmlElement('w:bookmarkStart')
        tag_start.set(qn('w:id'), '0')
        tag_start.set(qn('w:name'), bookmark_id)
        run._element.addprevious(tag_start)

        tag_end = OxmlElement('w:bookmarkEnd')
        tag_end.set(qn('w:id'), '0')
        tag_end.set(qn('w:name'), bookmark_id)
        run._element.addnext(tag_end)

    def _add_run(self, para, run_data: dict):
        """添加run到段落"""
        text = run_data.get("text", "")
        if not text:
            return

        r = para.add_run(text)

        if run_data.get("bold"):
            r.bold = True
        if run_data.get("italic"):
            r.italic = True
        if run_data.get("underline"):
            r.underline = True
        if run_data.get("superscript"):
            r.font.superscript = True
        if run_data.get("subscript"):
            r.font.subscript = True
        if run_data.get("color"):
            try:
                color = run_data["color"]
                if isinstance(color, str) and color.startswith("#"):
                    color = color[1:]
                r.font.color.rgb = RGBColor.from_string(color)
            except:
                pass
        if run_data.get("highlight"):
            try:
                highlight_color = run_data["highlight"]
                if isinstance(highlight_color, str) and highlight_color.startswith("#"):
                    highlight_color = highlight_color[1:]
                rPr = r._element.get_or_add_rPr()
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), highlight_color)
                rPr.append(shading)
            except:
                pass
