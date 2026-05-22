"""
目录块处理器 - 插入 TOC 域
"""
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .base import BlockHandler
from ..parser import TocBlock


class TocHandler(BlockHandler):
    """目录块处理器"""

    def can_handle(self, block) -> bool:
        return isinstance(block, TocBlock)

    def handle(self, block: TocBlock):
        """处理目录块 - 插入 TOC 域和提示文字"""
        # 获取各级别样式

        toc_style1 = styles["TOC1"]
        toc_style2 = styles["TOC2"]
        toc_style3 = styles["TOC3"]

        style_level_one = block.style_level_one or "TOC1"
        style_level_two = block.style_level_two or "TOC2"
        style_level_three = block.style_level_three or "TOC3"

        print(block.style_level_one, block.style_level_two, block.style_level_three)

        para = self.doc.add_paragraph()

        # 构建 TOC 域代码
        levels = block.levels or [1, 2, 3]
        # 将 levels 转为 Word 的 o/p 参数格式
        # 例如 [1,2,3] -> "1-3"
        if len(levels) >= 2:
            level_range = f"{min(levels)}-{max(levels)}"
        else:
            level_range = str(levels[0]) if levels else "1-3"

        # 创建 fldChar begin
        run_begin = para.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run_begin._element.append(fld_char_begin)

        # 创建 instrText（域代码）
        run_instr = para.add_run()
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        # 添加 \t 参数指定各级别样式
        toc_code = f' TOC \\o "{level_range}" \\h \\z \\u '
        toc_code += f'\\t "{style_level_one},{style_level_two},{style_level_three}"'
        instr_text.text = toc_code
        run_instr._element.append(instr_text)

        # 创建 fldChar separate
        run_sep = para.add_run()
        fld_char_sep = OxmlElement('w:fldChar')
        fld_char_sep.set(qn('w:fldCharType'), 'separate')
        run_sep._element.append(fld_char_sep)

        # 添加提示文字（用户在 Word 中更新域后会替换）
        run_hint = para.add_run("点击更新域来更新目录")

        # 创建 fldChar end
        run_end = para.add_run()
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run_end._element.append(fld_char_end)

        return para
