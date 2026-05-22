import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class JSON2DOCXConverter:
    def __init__(self, style_path):
        self.default_style_names = {
            "text": "default_style_text",
            "image": "default_style_image",
            "table": "default_style_table"
        }
        default_style_path = Path(__file__).parent / "default_style.json"
        self.default_styles = self._load_styles(default_style_path)
        user_styles = self._load_styles(style_path)
        self.styles = self._merge_styles(self.default_styles, user_styles)
        self.doc = Document()
        self._register_styles()

    def _load_styles(self, style_path):
        if not style_path:
            return {}
        style_path = Path(style_path)
        if not style_path.exists():
            return {}
        with open(style_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    def _merge_styles(self, base_styles, override_styles):
        styles = {name: value.copy() for name, value in base_styles.items()}
        for name, override in override_styles.items():
            if name in styles and isinstance(styles[name], dict) and isinstance(override, dict):
                merged = styles[name].copy()
                merged.update(override)
                styles[name] = merged
            else:
                styles[name] = override
        return styles

    def _register_styles(self):
        for tag, style in self.styles.items():
            style_name = tag
            # 检查样式是否已存在
            if style_name in [s.name for s in self.doc.styles]:
                # 如果存在，获取已有的样式对象
                para_style = self.doc.styles[style_name]
            else:
                # 如果不存在，创建新样式
                para_style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            # 设置字体
            font_name = style.get("font_name", "Times New Roman")
            font_name_east_asia = style.get("font_name_east_asia", "宋体")
            para_style.font.name = font_name
            para_style.font.name_east_asia = font_name_east_asia
            rFonts = para_style.font.element.rPr.rFonts
            rFonts.set(qn('w:eastAsia'), font_name_east_asia)
            # 设置字体大小
            font_size = style.get("font_size", 12)
            para_style.font.size = Pt(font_size)
            # 设置加粗
            bold = style.get("bold", False)
            para_style.font.bold = bold
            # 设置对齐方式
            alignment = style.get("alignment", "left")
            para_style.paragraph_format.alignment = self._get_alignment(alignment)
            # 设置字体颜色
            para_style.font.color.rgb = RGBColor(0, 0, 0)

            # 设置段前间距
            if style.get("space_before"):
                space_before = style["space_before"]
                units = space_before.get("units", "pt")
                value = space_before.get("value", 0)
                if units == "line":
                    para_style.paragraph_format.space_before = Pt(value * 12)  # 1行≈12磅
                else:
                    para_style.paragraph_format.space_before = Pt(value)
            else:
                para_style.paragraph_format.space_before = Pt(0)

            # 设置段后间距
            if style.get("space_after"):
                space_after = style["space_after"]
                units = space_after.get("units", "pt")
                value = space_after.get("value", 0)
                if units == "line":
                    para_style.paragraph_format.space_after = Pt(value * 12)  # 1行≈12磅
                else:
                    para_style.paragraph_format.space_after = Pt(value)
            else:
                para_style.paragraph_format.space_after = Pt(0)

            # 设置首行缩进
            if style.get("firstLineChars"):
                pPr = para_style.element.get_or_add_pPr()
                ind = OxmlElement('w:ind')
                ind.set(qn('w:firstLineChars'), str(int(style["firstLineChars"])))
                pPr.append(ind)

            # 设置行高
            if style.get("line_spacing"):
                line_spacing = style["line_spacing"]
                units = line_spacing.get("units", "line")
                value = line_spacing.get("value", 1.0)
                if units == "pt":
                    para_style.paragraph_format.line_spacing = Pt(value)
                    para_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                else:
                    para_style.paragraph_format.line_spacing = value

    def convert(self, json_path, output_path=None):
        self.json_dir = str(Path(json_path).parent.resolve())
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        blocks = data.get("blocks", [])
        for block in blocks:
            self._process_block(block)

        if output_path is None:
            output_path = str(Path(json_path).with_suffix('.docx'))
        self.doc.save(output_path)
        print(f"文档已生成: {output_path}")

    def _process_block(self, block):
        block_type = block.get("type")
        style_name = block.get("style", "")

        if block_type == "page-break":
            self._add_page_break()
            return

        if block_type == "text":
            self._add_text_block(block, style_name)
            return

        if block_type == "image":
            self._add_image_block(block, style_name)
            return

        if block_type == "table":
            self._add_table_block(block, style_name)
            return

        print(f"警告: 不支持的块类型 '{block_type}'，已跳过")

    def _add_text_block(self, block, style_name):
        # 获取样式
        default_style_name = self.default_style_names["text"]
        actual_style_name = style_name if style_name in self.styles else default_style_name

        para = self.doc.add_paragraph(style=actual_style_name)

        # 处理 runs
        runs = block.get("runs", [])
        if not runs:
            # 如果没有 runs，使用 block 的 text 字段（向后兼容）
            text = block.get("text", "")
            if text:
                para.add_run(text)
        else:
            for run in runs:
                self._add_run_to_para(para, run)

        return para

    def _add_run_to_para(self, para, run):
        text = run.get("text", "")
        if not text:
            return

        r = para.add_run(text)

        # 应用 inline 样式
        if run.get("bold"):
            r.bold = True
        if run.get("italic"):
            r.italic = True
        if run.get("underline"):
            r.underline = True
        if run.get("superscript"):
            r.font.superscript = True
        if run.get("subscript"):
            r.font.subscript = True
        if run.get("color"):
            try:
                color = run["color"]
                if isinstance(color, str) and color.startswith("#"):
                    color = color[1:]
                r.font.color.rgb = RGBColor.from_string(color)
            except:
                pass
        if run.get("highlight"):
            # 高亮背景色
            try:
                highlight_color = run["highlight"]
                if isinstance(highlight_color, str) and highlight_color.startswith("#"):
                    highlight_color = highlight_color[1:]
                # 使用 shading 设置背景色
                rPr = r._element.get_or_add_rPr()
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), highlight_color)
                rPr.append(shading)
            except:
                pass

    def _add_page_break(self):
        para = self.doc.add_paragraph()
        para.add_run().add_break(WD_BREAK.PAGE)
        return para

    def _add_image_block(self, block, style_name):
        # 获取样式
        default_style_name = self.default_style_names["image"]
        actual_style_name = style_name if style_name in self.styles else default_style_name
        style = self.styles.get(actual_style_name, self.styles.get(default_style_name, {}))

        para = self.doc.add_paragraph()
        para.alignment = self._get_alignment(style.get("alignment", "center"))

        if style.get("space_before"):
            space_before = style["space_before"]
            units = space_before.get("units", "pt")
            value = space_before.get("value", 0)
            if units == "line":
                para.paragraph_format.space_before = Pt(value * 12)
            else:
                para.paragraph_format.space_before = Pt(value)
        if style.get("space_after"):
            space_after = style["space_after"]
            units = space_after.get("units", "pt")
            value = space_after.get("value", 0)
            if units == "line":
                para.paragraph_format.space_after = Pt(value * 12)
            else:
                para.paragraph_format.space_after = Pt(value)

        src = block.get("src", "")
        if not src:
            para.add_run("[图片路径为空]")
            return

        img_path = src
        if not Path(src).is_absolute():
            img_path = str(Path(self.json_dir) / src)

        try:
            width_cm = style.get("width_cm", 10)
            para.add_run().add_picture(img_path, width=Cm(width_cm))
        except Exception as e:
            para.add_run(f"[图片加载失败: {src}]")

    def _set_table_no_border(self, table):
        tbl = table._element
        tblPr = tbl.tblPr

        tblBorders = OxmlElement('w:tblBorders')

        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), 'nil')
            tblBorders.append(element)

        tblPr.append(tblBorders)

    def _set_cell_border(self, cell, **kwargs):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')

        for edge, val in kwargs.items():
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), val)
            element.set(qn('w:sz'), "12")
            element.set(qn('w:space'), "0")
            element.set(qn('w:color'), "000000")
            tcBorders.append(element)

        tcPr.append(tcBorders)

    def _apply_table_borders(self, table, border_style):
        self._set_table_no_border(table)

        if border_style == "none":
            return

        if border_style == "grid":
            for row in table.rows:
                for cell in row.cells:
                    self._set_cell_border(
                        cell,
                        top="single",
                        bottom="single",
                        left="single",
                        right="single"
                    )
            return

        # three_line style
        rows = len(table.rows)
        if rows >= 1:
            for cell in table.rows[0].cells:
                self._set_cell_border(cell, top="single", bottom="single", left="nil", right="nil")

            for i in range(1, rows):
                for cell in table.rows[i].cells:
                    self._set_cell_border(cell, left="nil", right="nil")

            if rows >= 2:
                for cell in table.rows[-1].cells:
                    self._set_cell_border(cell, bottom="single", left="nil", right="nil")

    def _add_table_block(self, block, style_name):
        # 表格样式现在从 block 中获取，而不是从 style.json
        rows = block.get("rows", 0)
        cols = block.get("cols", 0)
        cells = block.get("cells", [])

        # 获取表格样式配置（从 block 中）
        header_style = block.get("header_style", "default_style_text")
        body_style = block.get("body_style", "default_style_text")
        border_style = block.get("border", "three_line")
        space_after_config = block.get("space_after", {"units": "pt", "value": 12})

        if rows <= 0 or cols <= 0:
            return

        table = self.doc.add_table(rows=rows, cols=cols)

        for i in range(min(rows, len(cells))):
            row_cells = cells[i]
            for j in range(min(cols, len(row_cells))):
                cell = table.cell(i, j)
                cell.text = str(row_cells[j])
                cell_style_name = header_style if i == 0 else body_style
                # 确保样式存在
                if cell_style_name not in self.styles:
                    cell_style_name = "default_style_text"
                for para in cell.paragraphs:
                    para.style = cell_style_name
                    para.paragraph_format.left_indent = 0
                    para.paragraph_format.first_line_indent = 0
                    para.paragraph_format.hanging_indent = 0

        self._apply_table_borders(table, border_style)

        # 表格后的间距
        para = self.doc.add_paragraph()
        if space_after_config:
            units = space_after_config.get("units", "pt")
            value = space_after_config.get("value", 12)
            if units == "line":
                para.paragraph_format.space_after = Pt(value * 12)
            else:
                para.paragraph_format.space_after = Pt(value)

    def _get_alignment(self, align_str):
        mapping = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        return mapping.get(align_str, WD_ALIGN_PARAGRAPH.LEFT)


def main():
    if len(sys.argv) < 2:
        print("用法: python converter.py <json文件路径> [输出docx路径] [style.json路径]")
        print("示例: python converter.py input.json output.docx my_style.json")
        return

    json_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    style_path = sys.argv[3] if len(sys.argv) > 3 else str(Path(__file__).parent / "style.json")

    converter = JSON2DOCXConverter(style_path)
    converter.convert(json_path, output_path)


if __name__ == "__main__":
    main()
